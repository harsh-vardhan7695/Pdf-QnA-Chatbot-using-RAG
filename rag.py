from langchain_community.document_loaders.web_base import WebBaseLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from pinecone import Pinecone, ServerlessSpec, PodSpec  
from langchain_pinecone import PineconeVectorStore
from langchain import hub
from langchain_groq import ChatGroq
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from dotenv import load_dotenv, find_dotenv
from nemoguardrails import LLMRails, RailsConfig
from nemoguardrails.integrations.langchain.runnable_rails import RunnableRails
import bs4
import os

from io import BytesIO
import time

from langchain_community.document_loaders import PyPDFLoader
from langchain_community.document_loaders import Docx2txtLoader
from langchain_community.document_loaders import TextLoader
from PyPDF2 import PdfReader
import docx
from langchain.docstore.document import Document

import asyncio

# from langchain.embeddings import SentenceTransformerEmbeddings
from langchain_huggingface import HuggingFaceEmbeddings
load_dotenv(find_dotenv())

os.environ["LANGCHAIN_TRACING_V2"] = "true"
os.environ["LANGCHAIN_API_KEY"] = os.getenv("LANGSMITH_API_KEY")
os.environ['LANGCHAIN_PROJECT'] = 'end-to-end-rag'

class RAG:

    def __init__(self, uploaded_files):
        
        self.vectorstore_index_name = "end-to-end-rag"
        self.uploaded_files = uploaded_files
        self.documents = []
        self.load_documents(uploaded_files)
        # self.embeddings = OpenAIEmbeddings(
        #     api_key=os.getenv("OPENAI_API_KEY"), model="text-embedding-3-small"
        # )
        self.embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        self.groq_llm = ChatGroq(
            api_key=os.getenv("GROQ_API_KEY"), 
            model="llama3-70b-8192", 
            temperature=0
        )
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000, 
            chunk_overlap=100
        )
        self.create_pinecone_index(self.vectorstore_index_name,384)
        self.vectorstore = PineconeVectorStore(
            index_name=self.vectorstore_index_name,
            embedding=self.embeddings,
            pinecone_api_key=os.getenv("PINECONE_API_KEY")
        )
        self.rag_prompt = hub.pull(
            "rlm/rag-prompt", 
            api_key=os.getenv("LANGSMITH_API_KEY")
        )
        config = RailsConfig.from_path("./config")

        self.guardrails = RunnableRails(config=config,llm=self.groq_llm)
        print("RAG initialization completed.")

    def load_documents(self, uploaded_files):
        for uploaded_file in uploaded_files:
            if uploaded_file is not None:
                bytes_data = uploaded_file.read()
                if uploaded_file.type == "application/pdf":
                    text = self.extract_text_from_pdf(BytesIO(bytes_data))
                elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                    text = self.extract_text_from_docx(BytesIO(bytes_data))
                elif uploaded_file.type == "text/plain":
                    text = self.extract_text_from_txt(BytesIO(bytes_data))
                else:
                    raise ValueError("Unsupported file type")
                document = Document(page_content=text)
                self.documents.append(document)
        print(f"Loaded {len(self.documents)} documents.")

    def extract_text_from_pdf(self, file_data):
        reader = PdfReader(file_data)
        text = ''
        for page_num in range(len(reader.pages)):
            text += reader.pages[page_num].extract_text()
        print("Extracted text from PDF.")
        return text

    def extract_text_from_docx(self, file_data):
        doc = docx.Document(file_data)
        text = '\n'.join([para.text for para in doc.paragraphs])
        print("Extracted text from DOCX.")
        return text

    def extract_text_from_txt(self, file_data):
        text = file_data.read().decode('utf-8')
        print("Extracted text from TXT.")
        return text

    # def load_documents(self):
    #     for folder in self.docs_folder:
    #         for file in os.listdir(folder):
    #             if file.endswith(".pdf"):
    #                 pdf_path = os.path.join(self.folder, file)
    #                 loader = PyPDFLoader(pdf_path)
    #                 self.documents = loader.load()  # Assign directly instead of extend
    #             elif file.endswith('.docx') or file.endswith('.doc'):
    #                 doc_path = os.path.join(self.folder, file)
    #                 loader = Docx2txtLoader(doc_path)
    #                 self.documents = loader.load()  # Assign directly instead of extend
    #             elif file.endswith('.txt'):
    #                 text_path = os.path.join(self.folder, file)
    #                 loader = TextLoader(text_path)
    #                 self.documents = loader.load()  # Assign directly instead of extend

    def create_pinecone_index(self, vectorstore_index_name,dimension):
        pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))  
        spec = ServerlessSpec(cloud='aws', region='us-east-1')  
        if vectorstore_index_name in pc.list_indexes().names():  
            pc.delete_index(vectorstore_index_name)  
        pc.create_index(  
            vectorstore_index_name,  
            dimension=dimension,
            metric='dotproduct',  
            spec=spec  
        )  
        while not pc.describe_index(vectorstore_index_name).status['ready']:  
            time.sleep(1)  
        print(f"Pinecone index {vectorstore_index_name} created and ready.")

    def load_docs_into_vectorstore_chain(self):
        # docs = self.loader.load()
        print("Splitting documents into chunks...")
        split_docs = self.text_splitter.split_documents(self.documents)
        print(f"Split documents into {len(split_docs)} chunks.")
        self.vectorstore.add_documents(split_docs)
        print("Documents loaded into vector store.")

    def format_docs(self, docs):
        return "\n\n".join(doc.page_content for doc in docs)
    
    def create_retrieval_chain(self):
        self.load_docs_into_vectorstore_chain()
        self.retriever = self.vectorstore.as_retriever()
        self.rag_chain = (
                        {
                            "context": self.retriever | self.format_docs, "question": RunnablePassthrough()
                        }
                        | self.rag_prompt
                        | self.groq_llm
                        | StrOutputParser()
                    )
        self.rag_chain = self.guardrails | self.rag_chain
        print("Retrieval chain created.")

    def qa(self, query, vectorstore_created):
        print(f"Received query: {query}")
        if not vectorstore_created:
            print("Creating retrieval chain...")
            self.create_retrieval_chain()
            vectorstore_created = True
        print("Invoking RAG chain...")
        response = self.rag_chain.invoke(query)
        print(f"Response: {response}")
        return response, vectorstore_created
    # def qa(self, query, vectorstore_created):
    #     if vectorstore_created:
    #         pass
    #     else:
    #         self.create_retrieval_chain(vectorstore_created)
    #     return self.rag_chain.invoke(query), True


    